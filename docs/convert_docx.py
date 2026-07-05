import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Set background color of a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tc_pr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner margins (padding) of a table cell in twentieths of a point (dxa)."""
    tc_pr = cell._element.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tc_mar.append(node)
    tc_pr.append(tc_mar)

def set_table_borders(table):
    """Apply clean light borders to a table."""
    tbl_pr = table._element.xpath('w:tblPr')
    if tbl_pr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>\n'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>\n'
            f'  <w:left w:val="none"/>\n'
            f'  <w:right w:val="none"/>\n'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>\n'
            f'  <w:insideV w:val="none"/>\n'
            f'</w:tblBorders>'
        )
        tbl_pr[0].append(borders)

def make_callout(doc, text):
    """Create a shaded gray box for code/diagram syntax."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    # Set left border to accent color
    tc_pr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="none"/>\n'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="52B788"/>\n'
        f'  <w:bottom w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'</w:tcBorders>'
    )
    tc_pr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    doc.add_paragraph().paragraph_format.space_before = Pt(6)

def build_docx():
    doc = Document()
    
    # Set Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Global Font settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55) # text-slate-700
    
    # Read markdown
    with open('tugas_8_perencanaan.md', 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_mermaid = False
    mermaid_lines = []
    in_table = False
    table_rows = []
    
    for line in lines:
        stripped = line.strip()
        
        # Handle Mermaid code blocks
        if stripped.startswith("```mermaid"):
            in_mermaid = True
            mermaid_lines = []
            continue
        elif in_mermaid and stripped.startswith("```"):
            in_mermaid = False
            # Render Mermaid block as a nice callout box
            make_callout(doc, "\n".join(mermaid_lines))
            continue
        elif in_mermaid:
            mermaid_lines.append(line.rstrip('\n'))
            continue
            
        # Handle Tables
        if stripped.startswith("|") and not in_mermaid:
            # Skip separator lines like |---|---|
            if re.match(r'^\|[\s:-|]+$', stripped):
                continue
            in_table = True
            # Parse row
            cols = [c.strip() for c in stripped.split("|")[1:-1]]
            table_rows.append(cols)
            continue
        elif in_table and not stripped.startswith("|"):
            in_table = False
            # Render collected table
            if table_rows:
                num_cols = len(table_rows[0])
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(table)
                
                for r_idx, row_data in enumerate(table_rows):
                    for c_idx, val in enumerate(row_data):
                        cell = table.cell(r_idx, c_idx)
                        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
                        
                        # Style headers differently
                        if r_idx == 0:
                            set_cell_background(cell, "F1F5F9") # Slate 100 background
                            
                        # Clean Markdown bold inside cells
                        clean_val = val.replace("**", "")
                        
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(4)
                        p.paragraph_format.space_after = Pt(4)
                        
                        run = p.add_run(clean_val)
                        run.font.size = Pt(9)
                        if r_idx == 0:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # dark slate
                doc.add_paragraph() # Add spacing after table
            table_rows = []
            
        if in_table:
            continue
            
        # Headers
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(18)
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        elif stripped.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2D, 0x6A, 0x4F) # Emerald Dark
        elif stripped.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(0x1B, 0x43, 0x32) # Darker green
        # Bullet Lists
        elif stripped.startswith("* ") or stripped.startswith("- "):
            clean_text = stripped[2:]
            
            # Check for nested bold text, e.g. * **Nama Tim:** AgriTernak
            bold_prefix = ""
            normal_text = clean_text
            match = re.match(r'^\*\*(.*?)\*\*(.*)', clean_text)
            if match:
                bold_prefix = match.group(1)
                normal_text = match.group(2)
                
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            
            if bold_prefix:
                b_run = p.add_run(bold_prefix)
                b_run.bold = True
                b_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            
            # Formatting links and bold in the rest of the text
            clean_rest = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', normal_text)
            p.add_run(clean_rest)
        # Image link
        elif stripped.startswith("![") and stripped.endswith(")"):
            match = re.search(r'\((.*?)\)', stripped)
            if match:
                img_path = match.group(1)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
                
                if os.path.exists(img_path):
                    p.add_run().add_picture(img_path, width=Inches(6.0))
                    # Add caption
                    caption_p = doc.add_paragraph()
                    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = caption_p.add_run("Gambar: Tangkapan Layar Trello Board AgriTernak Connect")
                    run.font.size = Pt(9)
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
                else:
                    run = p.add_run(f"\n[SISIPIKAN GAMBAR SCREENSHOT TRELLO '{img_path}' DI SINI]\n")
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26) # Red
        # Regular text
        elif stripped:
            # Skip page titles or subtitles parsed as headers already
            if stripped.startswith("**Mata Kuliah:**") or stripped.startswith("**Program Studi:**") or stripped.startswith("**Tugas:**"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(2)
                clean_line = stripped.replace("**", "")
                p.add_run(clean_line).italic = True
                continue
                
            if stripped == "---":
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run("__________________________________________________________________")
                run.font.color.rgb = RGBColor(0xE2, 0x8E, 0xF0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue
                
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            
            # Clean bold markers and links for plain text paragraph
            clean_text = stripped.replace("**", "")
            clean_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', clean_text)
            p.add_run(clean_text)
            
    doc.save('tugas_8_perencanaan.docx')
    print("Word document generated successfully!")

if __name__ == '__main__':
    build_docx()
